#!/usr/bin/env python3
"""Generate a Word document comparing the GW Migration Tool with Monday.com's Cross-Account Copier."""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
import os

def set_cell_shading(cell, color_hex):
    """Set cell background color."""
    shading = cell._element.get_or_add_tcPr()
    shd = shading.makeelement(qn('w:shd'), {
        qn('w:fill'): color_hex,
        qn('w:val'): 'clear',
    })
    # Remove existing shading
    for existing in shading.findall(qn('w:shd')):
        shading.remove(existing)
    shading.append(shd)

def set_cell_shading_v2(cell, color_hex):
    """Set cell background color (compatible approach)."""
    from docx.oxml import OxmlElement
    tc = cell._element
    tcPr = tc.find(qn('w:tcPr'))
    if tcPr is None:
        tcPr = OxmlElement('w:tcPr')
        tc.insert(0, tcPr)
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), color_hex)
    shd.set(qn('w:val'), 'clear')
    # Remove existing
    for existing in tcPr.findall(qn('w:shd')):
        tcPr.remove(existing)
    tcPr.append(shd)

def add_cell_text(cell, text, bold=False, color=None, size=10):
    """Add formatted text to a cell."""
    cell.text = ''
    p = cell.paragraphs[0]
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor(*color)

def create_comparison_doc():
    doc = Document()

    # Title
    title = doc.add_heading('Cross-Account Migration Tool Comparison', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Subtitle
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run('GW Workspace Migration Tool vs. Monday.com Cross-Account Copier')
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(100, 100, 100)

    doc.add_paragraph()  # spacer

    # Intro paragraph
    intro = doc.add_paragraph()
    run = intro.add_run(
        'This document compares the cross-account migration capabilities of the '
        'GW Workspace Migration Tool (custom-built Google Apps Script application) '
        'with Monday.com\'s native Cross-Account Copier feature. '
        'Only cross-account features are compared.'
    )
    run.font.size = Pt(11)

    doc.add_paragraph()  # spacer

    # === MAIN COMPARISON TABLE ===
    doc.add_heading('Feature Comparison', level=1)

    # Define comparison data
    rows_data = [
        # (Feature, GW Tool, Monday Copier, Category)
        ('SCOPE & STRUCTURE', '', '', 'header'),
        ('Migration Unit', 'Full workspace (all boards, groups, items)', 'Folder-based (boards must be in a folder)', ''),
        ('Board Types Supported', 'Main boards, Private boards, Shareable boards', 'Main boards only', ''),
        ('Dashboard Migration', 'Filtered out automatically (avoids duplicates)', 'Dashboards in folders are copied', ''),
        ('Workspace Creation', 'Automatically creates target workspace', 'Manual — must pre-create workspace', ''),
        ('Board Selection', 'Migrate all or select specific boards', 'All boards in the folder', ''),

        ('DATA MIGRATION', '', '', 'header'),
        ('Items & Column Values', 'Full migration with value mapping', 'Copies with board data', ''),
        ('Subitems', 'Full subitem migration with all column values', 'Included as part of board data', ''),
        ('Groups', 'Recreated with original structure and order', 'Included as part of board data', ''),
        ('Updates / Comments', 'Migrated via API with author attribution', 'Transfers if subscribers share email on both accounts', ''),
        ('Files / Attachments', 'Downloaded and re-uploaded to target account', 'Not explicitly guaranteed cross-account', ''),
        ('People Columns', 'Post-migration population with email-based user ID mapping', 'NOT transferred — people assignments are lost', ''),
        ('Status Columns', 'Full label and index mapping', 'Copies with board data', ''),
        ('Date Columns', 'Parsed and remapped correctly', 'Copies with board data', ''),
        ('Connect Boards Columns', 'Handled during migration', 'Hidden if linked board not in same folder', ''),

        ('USER MANAGEMENT', '', '', 'header'),
        ('User Migration', 'Invite users/guests to target account by email', 'Manual — must invite users separately', ''),
        ('Guest Migration', 'Board-level guest invitations preserved', 'Not supported', ''),
        ('User ID Mapping', 'Automatic email-based mapping across accounts', 'Not available — people columns lost', ''),
        ('User Activation Validation', 'Validates users are active before people column population', 'Not applicable', ''),

        ('VALIDATION & SAFETY', '', '', 'header'),
        ('Pre-Migration Validation', 'Full structural comparison (boards, items, columns, groups)', 'None — copies without validation', ''),
        ('Cross-Account Validation', 'Validates target workspace against source with detailed report', 'None', ''),
        ('Dry Run Mode', 'Preview people column changes before applying', 'Not available', ''),
        ('Progress Tracking', 'Real-time progress bar with status updates', 'No progress indicator', ''),
        ('Error Reporting', 'Detailed per-board, per-item error logging', 'Basic success/failure notification', ''),

        ('REQUIREMENTS & RESTRICTIONS', '', '', 'header'),
        ('Admin Requirement', 'API key access (no admin role required)', 'Must be admin on BOTH accounts with same email', ''),
        ('Regional Server Support', 'Works across any regions (API-based)', 'Cannot transfer between different regional servers (EU/US)', ''),
        ('Authentication', 'Dual API keys (source + target)', 'Same admin email on both accounts', ''),
        ('Automation Transfer', 'Not transferred (by design — avoids conflicts)', 'Active automations transfer; inactive do not', ''),
        ('Integration Transfer', 'Not transferred', 'Active integrations transfer (require re-auth)', ''),

        ('WORKFLOW & UX', '', '', 'header'),
        ('User Interface', 'Multi-tab React UI with guided workflow (7 steps)', 'Built-in Monday.com admin panel', ''),
        ('Batch Operations', 'Multiple boards processed in sequence', 'One board at a time recommended for reliability', ''),
        ('Post-Migration Steps', 'Guided: validate → invite users → populate people columns', 'Manual post-migration cleanup', ''),
        ('Workspace Refresh', 'Auto-refreshes workspace list after migration', 'Manual refresh required', ''),
        ('Target Account Selection', 'Dropdown selector for multiple target accounts', 'Fixed source-to-destination via admin panel', ''),
    ]

    # Create table
    table = doc.add_table(rows=1, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'

    # Header row
    header_cells = table.rows[0].cells
    headers = ['Feature', 'GW Migration Tool', 'Monday.com Cross-Account Copier']
    header_color = '0056B3'

    for i, header_text in enumerate(headers):
        cell = header_cells[i]
        add_cell_text(cell, header_text, bold=True, color=(255, 255, 255), size=11)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_shading_v2(cell, header_color)

    # Data rows
    for feature, gw_tool, monday_copier, category in rows_data:
        row = table.add_row()
        cells = row.cells

        if category == 'header':
            # Section header row
            # Merge all cells for section header
            cells[0].merge(cells[2])
            add_cell_text(cells[0], feature, bold=True, color=(255, 255, 255), size=11)
            cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
            set_cell_shading_v2(cells[0], '333333')
        else:
            add_cell_text(cells[0], feature, bold=True, size=10)
            set_cell_shading_v2(cells[0], 'F0F0F0')

            add_cell_text(cells[1], gw_tool, size=10)
            add_cell_text(cells[2], monday_copier, size=10)

    # Set column widths
    for row in table.rows:
        for i, width in enumerate([Cm(5), Cm(7.5), Cm(7.5)]):
            row.cells[i].width = width

    doc.add_paragraph()  # spacer

    # === KEY ADVANTAGES TABLE ===
    doc.add_heading('Key Advantages Summary', level=1)

    adv_table = doc.add_table(rows=1, cols=2)
    adv_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    adv_table.style = 'Table Grid'

    # Header
    h1 = adv_table.rows[0].cells[0]
    h2 = adv_table.rows[0].cells[1]
    add_cell_text(h1, 'GW Migration Tool Advantages', bold=True, color=(255, 255, 255), size=11)
    add_cell_text(h2, 'Monday.com Copier Advantages', bold=True, color=(255, 255, 255), size=11)
    h1.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    h2.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_cell_shading_v2(h1, '198754')  # green
    set_cell_shading_v2(h2, '0D6EFD')  # blue

    gw_advantages = [
        'Cross-regional server support (EU ↔ US)',
        'People column population with user ID mapping',
        'User & guest invitation to target account',
        'Pre-migration validation with detailed reports',
        'Dry run mode for safe previewing',
        'Real-time progress tracking',
        'Supports Private and Shareable boards',
        'No admin role required (API key based)',
        'Automatic workspace creation',
        'Post-migration guided workflow',
        'File transfer across accounts',
        'Dashboard filtering to prevent duplicates',
    ]

    monday_advantages = [
        'Native Monday.com feature (no setup required)',
        'Transfers active automations',
        'Transfers active integrations',
        'No API key management needed',
        'Built into admin panel UI',
        'Transfers cross-board automations (within same folder)',
        '',
        '',
        '',
        '',
        '',
        '',
    ]

    for gw_adv, mon_adv in zip(gw_advantages, monday_advantages):
        row = adv_table.add_row()
        c1 = row.cells[0]
        c2 = row.cells[1]
        if gw_adv:
            add_cell_text(c1, f'✓  {gw_adv}', size=10, color=(25, 135, 84))
        if mon_adv:
            add_cell_text(c2, f'✓  {mon_adv}', size=10, color=(13, 110, 253))

    # Set widths
    for row in adv_table.rows:
        row.cells[0].width = Cm(10)
        row.cells[1].width = Cm(10)

    doc.add_paragraph()  # spacer

    # Footer
    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run('Generated March 2026 — Cross-Account Features Only')
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(150, 150, 150)
    run.italic = True

    # Save
    output_path = os.path.join(os.path.dirname(__file__), 'Cross_Account_Migration_Comparison.docx')
    doc.save(output_path)
    print(f'Document saved to: {output_path}')
    return output_path

if __name__ == '__main__':
    create_comparison_doc()
